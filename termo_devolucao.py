import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import datetime

def formatar_moeda(valor):
    return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def centralizar_celula(celula):
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "center")
    tcPr.append(vAlign)

def gerar_termo_devolucao(nome, lista_de_bens):
    bens = pd.DataFrame(lista_de_bens)
    if bens.empty:
        return None

    doc = Document('timbrado.docx')

    # Título
    p = doc.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TERMO DE DEVOLUÇÃO")
    run.bold = True
    run.italic = False
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Texto de introdução
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Pelo presente termo, eu, ").bold = False
    p.add_run(nome).bold = True
    p.add_run(", declaro que devolvi ao Setor de Patrimônio o(s) bem(ns) abaixo discriminado(s), que se encontrava(m) sob minha guarda e responsabilidade:")

    doc.add_paragraph()

    # Tabela de bens
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
    larguras_colunas = [0.5, 4.5, 10, 4]
    cabecalhos = ['Patrimônio', 'Descrição', 'Complemento', 'Valor Atual']
    hdr_cells = tabela.rows[0].cells

    for i, heading in enumerate(cabecalhos):
        hdr_cells[i].text = heading
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0]
        run.font.size = Pt(12)
        run.bold = True
        centralizar_celula(hdr_cells[i])
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w',
                str(int(larguras_colunas[i] * 567)))
        tcPr.append(tcW)

    for _, bem in bens.iterrows():
        row_cells = tabela.add_row().cells
        row_cells[0].text = str(bem['Número Bem'])
        row_cells[1].text = str(bem['Descrição'])
        row_cells[2].text = str(bem['Complemento'])
        row_cells[3].text = formatar_moeda(bem['Valor Atual'])

        for cell in row_cells:
            centralizar_celula(cell)

    # Linha total
    total = bens['Valor Atual'].sum()
    total_row = tabela.add_row().cells
    total_row[0].merge(total_row[2])
    total_row[0].text = "TOTAL"
    run = total_row[0].paragraphs[0].runs[0]
    run.font.size = Pt(12)
    run.bold = True
    centralizar_celula(total_row[0])
    total_row[3].text = formatar_moeda(total)
    centralizar_celula(total_row[3])

    doc.add_paragraph()

    # Data
    hoje = datetime.now()
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    data_formatada = f"Brasília (DF), {hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"
    p_data = doc.add_paragraph(data_formatada)
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph()

    # Assinaturas
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run(nome).bold = True
    p2 = doc.add_paragraph("Assinado eletronicamente via SEI")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p3 = doc.add_paragraph("Declaro que recebi o(s) bem(ns) acima especificado(s):")
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("ANTÔNIO RODRIGUES DE SOUSA JÚNIOR").bold = True
    p5 = doc.add_paragraph("Supervisor de Patrimônio")
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6 = doc.add_paragraph("Assinado eletronicamente via SEI")
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    nome_arquivo = f"Termo_Devolucao_{nome.replace(' ', '_')}.docx"
    doc.save(nome_arquivo)
    return nome_arquivo
