from operator import truediv

import pandas as pd
import locale
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_termos(filtrar_ccusto=None):

    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    # Caminho para o arquivo Excel
    diretorio_atual = os.path.dirname(__file__)  # Obtém o diretório onde o script está localizado
    nome_arquivo_excel = 'acervo.xlsx'  # Nome do arquivo Excel
    caminho_da_planilha = os.path.join(diretorio_atual, nome_arquivo_excel)

    # Caminho para o arquivo de modelo (papel timbrado)
    nome_arquivo_modelo = 'timbrado.docx'  # Certifique-se de que o nome do arquivo está correto
    caminho_do_modelo = os.path.join(diretorio_atual, nome_arquivo_modelo)

    # Lendo os dados dos bens e responsáveis
    df_acervo = pd.read_excel(caminho_da_planilha, sheet_name='acervo')
    df_responsaveis = pd.read_excel(caminho_da_planilha, sheet_name='responsavel')

    # Unindo os DataFrames pelo campo 'ccustos'
    df_completo = pd.merge(df_acervo, df_responsaveis, on='ccustos')

    if filtrar_ccusto:
        df_completo = df_completo[df_completo['ccustos'] == filtrar_ccusto]

    # Agrupando os dados unidos por 'ccustos'
    grupos = df_completo.groupby('ccustos')

    # Defina seu texto padrão do termo aqui
    texto_padrao = """
    Pelo presente termo, eu, {responsavel}, matrícula n.º {matricula}, {funcao} do(a) {ccustos} do CFC, declaro que os bens patrimoniais abaixo discriminados se encontram na localização sob a minha guarda e responsabilidade.
    Assumo TOTAL responsabilidade pelos referidos bens, comprometendo-me a informar o Setor de Patrimônio quanto a qualquer alteração e/ou irregularidade, bem como zelar pela guarda e bom uso do patrimônio público.
    Em caso de extravio ou dano a bem sob a minha responsabilidade, comprometo-me a ressarcir o CFC dos prejuízos causados.
    Observações:
    Em caso de perda ou roubo do bem, o responsável deverá registrar boletim de ocorrência policial e apresentar ao Setor de Patrimônio;
    Ao final do mandato, função ou designação, o responsável deverá devolver o bem, se for o caso.
    No caso de movimentação e transferência de bens entre as unidades administrativas, o Setor de Patrimônio utilizará o Termo de Transferência disponível no SEI, que será apensado a processo específico até a emissão de um novo termo atualizado.
    """

    # Para cada ccustos, criar um Termo de Responsabilidade
    for ccustos, grupo in grupos:
        responsavel = grupo.iloc[0]

        # Criar um novo documento Word
        documento = Document(caminho_do_modelo)
        cabecalho = documento.add_paragraph()
        cabecalho_run = cabecalho.add_run(f'Termo de Responsabilidade - {ccustos}')
        cabecalho_run.font.bold = True
        cabecalho.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cabecalho_run.font.size = Pt(16)

        # Calcular a soma dos valores
        soma_valores = sum(grupo['valor_atual'])

        # Ordenar a tabela por localização
        grupo_ordenado = grupo.sort_values(by='numero')

        # Formatar o texto padrão com informações do responsável e do ccustos
        for paragraph in texto_padrao.split('\n'):
            paragrafo = documento.add_paragraph()
            paragrafo.paragraph_format.first_line_indent = Inches(0.59)  # Recuo de 1,5 cm na primeira linha
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragrafo.add_run(paragraph.format(
                responsavel=responsavel['responsavel'],
                matricula=responsavel['matricula'],
                funcao=responsavel['funcao'],
                ccustos=responsavel['ccustos']
            ))

        # Adicionar uma tabela para os bens
        tabela = documento.add_table(rows=1, cols=5)
        tabela.style = 'Table Grid'
        for cell in tabela.columns[0].cells:
            cell.width = Inches(0.5)
        for cell in tabela.columns[1].cells:
            cell.width = Inches(0.5)
        for cell in tabela.columns[2].cells:
            cell.width = Inches(3.0)
        for cell in tabela.columns[3].cells:
            cell.width = Inches(1.0)
        for cell in tabela.columns[4].cells:
            cell.width = Inches(0.75)

        # Cabeçalhos da tabela
        hdr_cells = tabela.rows[0].cells
        for i, heading in enumerate(['Número Bem', 'Descrição', 'Complemento', 'Localização', 'Valor Atual']):
            hdr_cells[i].text = heading

            # Obter o parágrafo na célula do cabeçalho
            paragraph = hdr_cells[i].paragraphs[0]

            # Remover o run existente e adicionar um novo com o tamanho de fonte desejado
            paragraph.clear()
            run = paragraph.add_run(heading)
            run.font.bold = True
            run.font.size = Pt(10)  # Definir o tamanho da fonte desejado aqui
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Preencher a tabela com os dados dos bens
        for _, bem in grupo_ordenado.iterrows():
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(bem['numero'])
            row_cells[1].text = bem['descricao']
            row_cells[2].text = str(bem['complemento'])
            row_cells[3].text = str(bem['localizacao'])

            # Formatar e alinhar o valor atual à direita
            valor_formatado = locale.format_string("%1.2f", bem['valor_atual'], grouping=True)
            paragrafo_valor = row_cells[4].paragraphs[0]
            paragrafo_valor.clear()
            run_valor = paragrafo_valor.add_run(valor_formatado)
            run_valor.font.size = Pt(9)  # Configurar o tamanho da fonte para 9
            paragrafo_valor.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Configurar o mesmo tamanho de fonte para todas as células
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        # Adicionar linha final "TOTAL" com células mescladas e somatório na última coluna
        ultima_linha = tabela.add_row().cells
        ultima_linha[0].merge(ultima_linha[1]).merge(ultima_linha[2]).merge(ultima_linha[3])
        ultima_linha[0].text = "TOTAL"
        ultima_linha[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar o valor total na última coluna
        valor_total_formatado = locale.format_string("%.2f", soma_valores, grouping=True)
        ultima_linha[4].text = valor_total_formatado
        ultima_linha[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Ajuste de formatação para a última linha
        for cell in ultima_linha:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.bold = True

        # Adicionar a seção de assinatura
        documento.add_paragraph()
        paragrafo_assinatura = documento.add_paragraph()
        paragrafo_assinatura.add_run(f"{responsavel['responsavel']}\n{responsavel['funcao']} do(a) {ccustos} do CFC")
        paragrafo_assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragrafo_assinatura.runs:
            run.font.size = Pt(12)

        # Salvar o documento com um nome único baseado no ccustos
        documento.save(f'Termo_de_Responsabilidade_{ccustos}.docx')
        df_para_excel = grupo[['numero', 'descricao', 'complemento', 'localizacao', 'valor_atual']]
        nome_arquivo_excel = f'planilha_{ccustos}.xlsx'
        df_para_excel.to_excel(nome_arquivo_excel, index=False)

    return "Termos gerados com sucesso!"