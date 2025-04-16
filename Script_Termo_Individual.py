import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
import locale

# Configura o locale para o padrão brasileiro
locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

def centralizar_celula(celula):
    """Centraliza o texto horizontal e verticalmente em uma célula."""
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centraliza horizontalmente
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "center")
    tcPr.append(vAlign)

def criar_termo_responsabilidade(nome, bens):
    # Carregar o modelo de documento com o fundo padrão
    template_path = '../pythonProject1/Automacao/Termo_de_Responsabilidade_Individual/timbrado.docx'
    doc = Document(template_path)

    # Adiciona cabeçalho
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TERMO DE RESPONSABILIDADE')
    run.bold = True
    run.font.size = Pt(14)
    run.italic = False

    # Adiciona um espaço entre o título e o início do parágrafo
    doc.add_paragraph("")

    # Adiciona parágrafo inicial
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("Pelo presente termo, eu, ").bold = False
    p.add_run(f"{nome}").bold = True
    p.add_run(", declaro que o(s) equipamento(s) abaixo discriminado(s) se encontra(m) sob a minha guarda e responsabilidade.")

    # Adiciona um espaço antes da tabela
    doc.add_paragraph("")

    # Adiciona uma tabela para os bens
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Define as larguras das colunas em centímetros (exemplo: 3 cm, 5 cm, 7 cm, 4 cm)
    larguras_colunas = [0.5, 4.5, 10, 4]

    # Cabeçalhos da tabela
    hdr_cells = tabela.rows[0].cells
    for i, heading in enumerate(['Patrimônio', 'Descrição', 'Complemento', 'Valor Atual']):
        hdr_cells[i].text = heading
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.runs[0]
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.bold = True
        centralizar_celula(hdr_cells[i])  # Centraliza texto no cabeçalho

        # Define a largura da coluna
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w',
                str(int(larguras_colunas[i] * 567)))  # Converte cm para twips (1 cm = 567 twips)
        tcPr.append(tcW)

    # Adiciona os dados dos bens à tabela
    for _, bem in bens.iterrows():
        # Converte todos os valores para string, substituindo NaN por string vazia
        patrimonio = str(bem['Patrimônio']) if not pd.isna(bem['Patrimônio']) else ""
        descricao = str(bem['Descrição']) if not pd.isna(bem['Descrição']) else ""
        complemento = str(bem['Complemento']) if not pd.isna(bem['Complemento']) else ""
        valor_atual = (
            locale.currency(bem['Valor Atual'], grouping=True, symbol=False).strip()
            if not pd.isna(bem['Valor Atual'])
            else ""
        )

        row_cells = tabela.add_row().cells
        row_cells[0].text = patrimonio
        row_cells[1].text = descricao
        row_cells[2].text = complemento
        row_cells[3].text = valor_atual

        # Centraliza todas as células na linha atual
        for cell in row_cells:
            centralizar_celula(cell)

    # Adiciona linha com o valor total e formata a célula "TOTAL"
    valor_total = bens['Valor Atual'].sum()
    total_row = tabela.add_row().cells
    total_row[0].merge(total_row[2])  # Mescla as colunas A, B, C
    total_row[0].text = "TOTAL"
    run = total_row[0].paragraphs[0].runs[0]
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.bold = True
    centralizar_celula(total_row[0])  # Centraliza o texto "TOTAL"
    total_row[3].text = locale.currency(valor_total, grouping=True, symbol=False).strip()
    centralizar_celula(total_row[3])  # Centraliza o valor total

    # Substitui o texto do parágrafo adicional
    novo_paragrafo = """
    Comprometo-me a:

    1) zelar pela guarda, uso adequado e conservação do(s) bem(ns), utilizando-o(s) exclusivamente para fins profissionais do CFC;

    2) informar imediatamente ao Setor de Patrimônio qualquer dano, inutilização, perda ou roubo, apresentando boletim de ocorrência quando necessário;

    3) ressarcir o CFC por danos ou perdas decorrentes de negligência do responsável, após decisão da Câmara de Assuntos Administrativos (CAD) e homologação pelo Plenário do CFC, em conformidade com o Manual de Gestão Patrimonial do CFC;

    4) devolver o(s) equipamento(s) e acessórios ao término do vínculo, mediante solicitação ou em caso de substituição, em condições compatíveis com o uso; e

    5) fornecer informações sobre o(s) bem(ns) sempre que solicitado, especialmente durante o inventário patrimonial.
    """

    # Adiciona um espaço após a tabela
    doc.add_paragraph("")

    for paragraph in novo_paragrafo.strip().split('\n'):
        p = doc.add_paragraph(paragraph.strip())
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Adiciona um espaço após a tabela
    doc.add_paragraph("")

    # Adiciona parágrafo adicional
    doc.add_paragraph("Declaro estar ciente das responsabilidades mencionadas acima e assumo total responsabilidade pelos bens listados.")

    # Adiciona um espaço após a tabela
    doc.add_paragraph("")

    # Adiciona parte "Assinado eletronicamente via SEI"
    p_assinado = doc.add_paragraph()
    p_assinado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_assinado.add_run(nome).bold = True

    p_assinado = doc.add_paragraph()
    p_assinado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_assinado.add_run("Assinado eletronicamente via SEI").bold = False

    # Salva o documento com o nome do responsável
    doc.save(f"Termo_{nome.replace(' ', '_')}.docx")


if __name__ == "__main__":
    # Carregar os dados do Excel
    df_geral = pd.read_excel('geral.xlsx', sheet_name='dados')

    # Converte todas as colunas relevantes para string, garantindo que NaN seja tratado
    df_geral['Patrimônio'] = df_geral['Patrimônio'].astype(str)
    df_geral['Descrição'] = df_geral['Descrição'].astype(str)
    df_geral['Complemento'] = df_geral['Complemento'].astype(str)

    # Agrupar por nome e criar termos de responsabilidade
    for nome, grupo in df_geral.groupby('Nome'):
        criar_termo_responsabilidade(nome, grupo)

    print("Termos de responsabilidade criados com sucesso!")